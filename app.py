# from flask import Flask, request, jsonify, render_template
# import os
# import pdfplumber
# import spacy
# import re
# from docx import Document

# app = Flask(__name__)

# UPLOAD_FOLDER = "uploads"
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
# ALLOWED_EXTENSIONS = {"pdf", "docx"}

# nlp = spacy.load("en_core_web_sm")

# TECHNICAL_SKILLS = [
#     "Python", "Django", "Flask", "SQL", "AWS", "Docker", "JavaScript", "React",
#     "Node.js", "Machine Learning", "Deep Learning", "NLP", "TensorFlow", "PyTorch"
# ]

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text(filepath):
#     if filepath.endswith(".pdf"):
#         return extract_text_from_pdf(filepath)
#     elif filepath.endswith(".docx"):
#         return extract_text_from_docx(filepath)
#     return ""

# def extract_text_from_pdf(filepath):
#     text = ""
#     with pdfplumber.open(filepath) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text.strip() if text else "Text extraction failed. Try another format."

# def extract_text_from_docx(filepath):
#     text = ""
#     doc = Document(filepath)
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text.strip()

# def extract_education(text):
#     education_patterns = [r"\b(B\.?Tech|MCA|BCA|M\.?Tech|BSc|MSc|MBA|PhD|Bachelor|Master|Diploma)\b"]
#     matches = []
#     for pattern in education_patterns:
#         matches.extend(re.findall(pattern, text, re.IGNORECASE))
#     return list(set(matches))

# def extract_experience(text):
#     experience_pattern = r"(\d+)\+?\s*(years|yrs|year|exp)\s*(of experience|in)?"
#     matches = re.findall(experience_pattern, text, re.IGNORECASE)
#     return [f"{match[0]} years" for match in matches] if matches else []

# def extract_skills(text):
#     extracted_skills = [skill for skill in TECHNICAL_SKILLS if re.search(rf"\b{skill}\b", text, re.IGNORECASE)]
#     return list(set(extracted_skills))

# def process_resume_text(text):
#     skills = extract_skills(text)
#     education = extract_education(text)
#     experience = extract_experience(text)
#     return {"skills": skills, "education": education, "experience": experience}

# def calculate_match_score(resume_data, job_requirements):
#     missing_skills = list(set(job_requirements["skills"]) - set(resume_data["skills"]))
#     missing_education = list(set(job_requirements["education"]) - set(resume_data["education"]))
#     missing_experience = list(set(job_requirements["experience"]) - set(resume_data["experience"]))
    
#     total_requirements = sum(map(len, [job_requirements["skills"], job_requirements["education"], job_requirements["experience"]]))
#     total_match = total_requirements - sum(map(len, [missing_skills, missing_education, missing_experience]))
#     match_score = (total_match / total_requirements) * 100 if total_requirements > 0 else 0
    
#     return match_score, {"missing_skills": missing_skills, "missing_education": missing_education, "missing_experience": missing_experience}

# @app.route("/")
# def home():
#     return render_template("index.html")

# @app.route("/upload", methods=["POST"])
# def upload_resume():
#     if "file" not in request.files:
#         return jsonify({"error": "No file part"}), 400

#     file = request.files["file"]
#     if file.filename == "":
#         return jsonify({"error": "No selected file"}), 400

#     if file and allowed_file(file.filename):
#         filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
#         file.save(filepath)
        
#         extracted_text = extract_text(filepath)
#         if not extracted_text:
#             return jsonify({"error": "Could not extract text from the file"}), 400
        
#         extracted_data = process_resume_text(extracted_text)
#         return jsonify({"message": "File uploaded successfully", "resume_data": extracted_data})

#     return jsonify({"error": "Invalid file format"}), 400

# @app.route("/compare", methods=["POST"])
# def compare_resume():
#     data = request.json
#     if "resume_data" not in data or "job_description" not in data:
#         return jsonify({"error": "Missing resume data or job description"}), 400

#     job_requirements = process_resume_text(data["job_description"])
#     match_score, missing_data = calculate_match_score(data["resume_data"], job_requirements)

#     return jsonify({"message": "Comparison completed", "match_score": match_score, **missing_data})

# if __name__ == "__main__":
#     port = int(os.environ.get("PORT", 5000))
#     app.run(host="0.0.0.0", port=port, debug=True)


from flask import Flask, request, jsonify, render_template
import os
import pdfplumber
import spacy
import re
from docx import Document
from collections import defaultdict

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {"pdf", "docx"}

nlp = spacy.load("en_core_web_sm")

TECHNICAL_SKILLS = {
    "Python": ["Python", "Py"],
    "Django": ["Django"],
    "Flask": ["Flask"],
    "SQL": ["SQL", "MySQL", "PostgreSQL", "SQLite"],
    "AWS": ["AWS", "Amazon Web Services"],
    "Docker": ["Docker"],
    "JavaScript": ["JavaScript", "JS"],
    "React": ["React", "ReactJS"],
    "Node.js": ["Node.js", "Node"],
    "Machine Learning": ["Machine Learning", "ML"],
    "Deep Learning": ["Deep Learning", "DL"],
    "NLP": ["Natural Language Processing", "NLP"],
    "TensorFlow": ["TensorFlow", "TF"],
    "PyTorch": ["PyTorch"]
}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(filepath):
    text = ""
    if filepath.endswith(".pdf"):
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    elif filepath.endswith(".docx"):
        doc = Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text.strip()

def extract_education(text):
    education_patterns = [r"\b(B\.?Tech|MCA|BCA|M\.?Tech|BSc|MSc|MBA|PhD|Bachelor|Master|Diploma)\b"]
    matches = []
    for pattern in education_patterns:
        matches.extend(re.findall(pattern, text, re.IGNORECASE))
    return list(set(matches))

def extract_experience(text):
    experience_pattern = r"(\d+)\+?\s*(years|yrs|year)\s*(of experience)?"
    matches = re.findall(experience_pattern, text, re.IGNORECASE)
    return [f"{match[0]} years" for match in matches] if matches else []

def extract_skills(text):
    extracted_skills = set()
    for key, synonyms in TECHNICAL_SKILLS.items():
        for synonym in synonyms:
            if re.search(rf"\b{synonym}\b", text, re.IGNORECASE):
                extracted_skills.add(key)
    return list(extracted_skills)

def process_resume_text(text):
    return {
        "skills": extract_skills(text),
        "education": extract_education(text),
        "experience": extract_experience(text)
    }

def calculate_match_score(resume_data, job_requirements):
    matched_skills = set(resume_data["skills"]) & set(job_requirements["skills"])
    missing_skills = list(set(job_requirements["skills"]) - matched_skills)
    missing_education = list(set(job_requirements["education"]) - set(resume_data["education"]))
    missing_experience = list(set(job_requirements["experience"]) - set(resume_data["experience"]))
    
    total_requirements = sum(map(len, [job_requirements["skills"], job_requirements["education"], job_requirements["experience"]]))
    total_match = sum(map(len, [matched_skills, resume_data["education"], resume_data["experience"]]))
    
    match_score = (total_match / total_requirements) * 100 if total_requirements > 0 else 0
    
    return match_score, {"missing_skills": missing_skills, "missing_education": missing_education, "missing_experience": missing_experience}

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    if file and allowed_file(file.filename):
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(filepath)
        extracted_text = extract_text(filepath)
        extracted_data = process_resume_text(extracted_text)
        return jsonify({"message": "File uploaded successfully", "resume_data": extracted_data})
    
    return jsonify({"error": "Invalid file format"}), 400

@app.route("/compare", methods=["POST"])
def compare_resume():
    data = request.json
    if "resume_data" not in data or "job_description" not in data:
        return jsonify({"error": "Missing resume data or job description"}), 400
    
    resume_data = data["resume_data"]
    job_description = data["job_description"]
    job_requirements = process_resume_text(job_description)
    match_score, missing_data = calculate_match_score(resume_data, job_requirements)
    
    return jsonify({
        "message": "Comparison completed",
        "match_score": match_score,
        "missing_skills": missing_data["missing_skills"],
        "missing_education": missing_data["missing_education"],
        "missing_experience": missing_data["missing_experience"]
    })

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)